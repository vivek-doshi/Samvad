import logging
import uuid
from pathlib import Path
from datetime import datetime, timezone

import pdfplumber
import docx
import pandas as pd

from backend.rag.chunkers import SemanticChunker
from backend.rag.embedder import Embedder
from backend.rag.bm25_index import BM25Index
from backend.db.db_client import DBClient
from backend.security.document_sanitiser import DocumentSanitiser

logger = logging.getLogger(__name__)


class DocumentIngester:
    """Parse → sanitise → chunk → embed → store in ChromaDB and BM25."""

    def __init__(
        self,
        embedder: Embedder,
        chroma_client,
        bm25: BM25Index,
        upload_dir: str = "runtime/user_uploads",
    ):
        self.embedder = embedder
        self.chroma_client = chroma_client
        self.bm25 = bm25
        self.upload_dir = Path(upload_dir)
        self.sanitiser = DocumentSanitiser()

    async def ingest(
        self,
        file_path: str,
        filename: str,
        file_type: str,
        session_id: str,
        user_id: str,
        doc_id: str,
        db: DBClient,
    ) -> dict:
        # 1. Parse file
        ext = file_type.lower()
        if ext == "pdf":
            text = self._parse_pdf(file_path)
        elif ext == "docx":
            text = self._parse_docx(file_path)
        elif ext == "csv":
            text = self._parse_csv(file_path)
        elif ext == "xlsx":
            text = self._parse_xlsx(file_path)
        else:
            text = Path(file_path).read_text(encoding="utf-8")

        # 2. Sanitise
        san_result = self.sanitiser.sanitise(text)
        status = san_result.get("status", "clean")
        flags = san_result.get("flags", [])
        if status == "quarantined":
            raise ValueError(f"Document quarantined: {flags}")
        if status == "flagged":
            text = san_result.get("sanitised_text", text)

        # 3. Chunk
        chunker = SemanticChunker()
        chunks = chunker.chunk_document(text, filename, session_id, doc_id)

        # 4. Embed
        texts = [c.text for c in chunks]
        embeddings = self.embedder.embed_texts(texts)

        # 5. Store in ChromaDB
        collection_name = f"user_{doc_id}"
        col = self.chroma_client.get_or_create_collection(collection_name)
        batch_size = 500
        for i in range(0, len(chunks), batch_size):
            batch_chunks = chunks[i : i + batch_size]
            batch_embeddings = embeddings[i : i + batch_size]
            col.add(
                ids=[c.chunk_id for c in batch_chunks],
                embeddings=batch_embeddings,
                documents=[c.text for c in batch_chunks],
                metadatas=[{**c.metadata, "chunk_id": c.chunk_id,
                             "chunk_level": c.chunk_level,
                             "parent_chunk_id": c.parent_chunk_id or "",
                             "doc_type": c.doc_type,
                             "source_name": c.source_name}
                           for c in batch_chunks],
            )

        # 6. Build BM25
        self.bm25.build_for_collection(collection_name, chunks)

        # 7. Update DB
        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            """UPDATE user_documents
               SET chunk_count=?, chroma_collection=?,
                   sanitisation_status=?, indexed_at=?
               WHERE doc_id=?""",
            (len(chunks), collection_name, status, now, doc_id),
        )

        return {
            "chunk_count": len(chunks),
            "chroma_collection": collection_name,
            "sanitisation_status": status,
            "flags": flags,
        }

    def _parse_pdf(self, path: str) -> str:
        parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                parts.append(f"\n[Page {page_num}]\n")
                text = page.extract_text() or ""
                parts.append(text)
                tables = page.extract_tables() or []
                for table in tables:
                    rows = []
                    for row in table:
                        rows.append("|" + "|".join(str(cell or "") for cell in row) + "|")
                    parts.append("\n".join(rows))
        return "\n".join(parts)

    def _parse_docx(self, path: str) -> str:
        doc = docx.Document(path)
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("|" + "|".join(cell.text for cell in row.cells) + "|")
        return "\n".join(parts)

    def _parse_csv(self, path: str) -> str:
        df = pd.read_csv(path)
        parts = [
            f"Columns: {', '.join(df.columns.tolist())}",
            df.describe(include="all").to_string(),
        ]
        for start in range(0, len(df), 50):
            chunk = df.iloc[start : start + 50]
            parts.append(chunk.to_markdown(index=False))
        return "\n\n".join(parts)

    def _parse_xlsx(self, path: str) -> str:
        xl = pd.ExcelFile(path)
        parts: list[str] = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            parts.append(f"## Sheet: {sheet}\n")
            parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            parts.append(df.describe(include="all").to_string())
            for start in range(0, len(df), 50):
                chunk = df.iloc[start : start + 50]
                parts.append(chunk.to_markdown(index=False))
        return "\n\n".join(parts)
